#!/usr/bin/env python3
"""Structured parser runner for ArkLoop Book KB M1.1.

Contract:
  bookparser_runner.py --mime application/pdf --max-pages 0 < blob.bin > out.json

The JSON shape is:
  {"meta": {...}, "blocks": [{"type": "paragraph", "text": "...", ...}]}
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import re
import sys
from collections import Counter
from dataclasses import dataclass, field
from typing import Any


MIN_IMAGE_EDGE_PX = 32
MIN_IMAGE_BYTES = 4096
MAX_IMAGE_CANDIDATES_PER_PAGE = 50

MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
IMAGE_MIMES = {"image/png", "image/jpeg", "image/jpg", "image/webp"}


@dataclass
class Block:
    type: str
    text: str
    heading_path: list[str] = field(default_factory=list)
    heading_inferred: bool = False
    heading_confidence: float = 0.0
    metadata: dict[str, Any] = field(default_factory=dict)

    def as_json(self) -> dict[str, Any]:
        return {
            "type": self.type,
            "text": self.text,
            "heading_path": self.heading_path,
            "heading_inferred": self.heading_inferred,
            "heading_confidence": self.heading_confidence,
            "metadata": self.metadata,
        }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--mime", required=True)
    parser.add_argument("--max-pages", type=int, default=0)
    args = parser.parse_args()

    data = sys.stdin.buffer.read()
    mime = normalize_mime(args.mime)
    try:
        if mime == MIME_PDF:
            payload = parse_pdf(data, args.max_pages)
        elif mime == MIME_DOCX:
            payload = parse_docx(data)
        elif mime in IMAGE_MIMES:
            payload = parse_image(data, mime)
        else:
            raise ValueError(f"unsupported mime: {args.mime}")
    except Exception as exc:  # noqa: BLE001 - CLI boundary, stderr is diagnostic.
        print(f"bookparser_runner: {exc}", file=sys.stderr)
        return 2

    payload.setdefault("meta", {})
    payload["meta"]["source_mime"] = mime
    payload["meta"]["byte_size"] = len(data)
    finalize(payload)
    json.dump(payload, sys.stdout, ensure_ascii=False)
    sys.stdout.write("\n")
    return 0


def normalize_mime(raw: str) -> str:
    base = raw.split(";", 1)[0].strip().lower()
    if base == "image/jpg":
        return "image/jpeg"
    if base in {"text/x-docx", "application/docx"}:
        return MIME_DOCX
    return base


def parse_pdf(data: bytes, max_pages: int) -> dict[str, Any]:
    try:
        import fitz  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("PyMuPDF (fitz) is required for PDF parsing") from exc

    blocks: list[Block] = []
    meta: dict[str, Any] = {"parser": "pymupdf+pdfplumber+tesseract"}
    doc = fitz.open(stream=data, filetype="pdf")
    total_pages = len(doc)
    page_limit = total_pages if max_pages <= 0 else min(total_pages, max_pages)
    meta["page_count"] = total_pages
    meta["parsed_pages"] = page_limit

    text_lines_by_page: dict[int, list[str]] = {}
    line_infos: list[dict[str, Any]] = []
    for page_index in range(page_limit):
        page = doc[page_index]
        info = page.get_text("dict")
        page_lines: list[str] = []
        for text_block in info.get("blocks", []):
            if text_block.get("type") != 0:
                continue
            for line in text_block.get("lines", []):
                spans = [s for s in line.get("spans", []) if s.get("text", "").strip()]
                if not spans:
                    continue
                text = normalize_space("".join(s.get("text", "") for s in spans))
                if not text:
                    continue
                size = weighted_font_size(spans)
                font = dominant_font(spans)
                line_infos.append(
                    {
                        "text": text,
                        "page": page_index + 1,
                        "font_size": size,
                        "font": font,
                        "bbox": line.get("bbox"),
                    }
                )
                page_lines.append(text)
        text_lines_by_page[page_index + 1] = page_lines

    add_line_blocks(blocks, line_infos)
    add_pdf_tables(blocks, data, page_limit)
    add_pdf_images(blocks, doc, page_limit, text_lines_by_page, meta)

    if not any(b.text.strip() for b in blocks if b.type in {"paragraph", "heading"}):
        blocks.extend(ocr_pdf_pages(doc, page_limit))

    return {"meta": meta, "blocks": [b.as_json() for b in blocks]}


def add_line_blocks(blocks: list[Block], line_infos: list[dict[str, Any]]) -> None:
    if not line_infos:
        return
    body_size = median([float(x["font_size"]) for x in line_infos])
    heading_sizes = sorted(
        {round(float(x["font_size"])) for x in line_infos if is_heading_line(x, body_size)},
        reverse=True,
    )
    levels = {size: idx + 1 for idx, size in enumerate(heading_sizes[:6])}
    heading_stack: list[str] = []
    paragraph_buf: list[str] = []
    paragraph_meta: dict[str, Any] = {}

    def flush() -> None:
        nonlocal paragraph_buf, paragraph_meta
        text = normalize_space(" ".join(paragraph_buf))
        if text:
            blocks.append(Block(type="paragraph", text=text, heading_path=list(heading_stack), metadata=dict(paragraph_meta)))
        paragraph_buf = []
        paragraph_meta = {}

    for line in line_infos:
        text = str(line["text"])
        font_size = round(float(line["font_size"]))
        if is_heading_line(line, body_size) and font_size in levels:
            flush()
            level = levels[font_size]
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(text)
            blocks.append(
                Block(
                    type="heading",
                    text=text,
                    heading_path=list(heading_stack),
                    heading_inferred=True,
                    heading_confidence=heading_confidence(line, body_size),
                    metadata={"page": line["page"], "font_size": line["font_size"], "font": line["font"]},
                )
            )
            continue
        if looks_like_formula(text):
            flush()
            blocks.append(Block(type="formula", text=text, heading_path=list(heading_stack), metadata={"page": line["page"]}))
            continue
        paragraph_buf.append(text)
        paragraph_meta.setdefault("page", line["page"])
    flush()


def add_pdf_tables(blocks: list[Block], data: bytes, page_limit: int) -> None:
    try:
        import pdfplumber  # type: ignore
    except Exception:
        return
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page_index, page in enumerate(pdf.pages[:page_limit]):
                for table_index, table in enumerate(page.extract_tables() or []):
                    md = table_to_markdown(table)
                    if md:
                        blocks.append(
                            Block(
                                type="table",
                                text=md,
                                metadata={"page": page_index + 1, "table_index": table_index},
                            )
                        )
    except Exception:
        return


def add_pdf_images(blocks: list[Block], doc: Any, page_limit: int, text_lines_by_page: dict[int, list[str]], meta: dict[str, Any]) -> None:
    skipped = 0
    for page_index in range(page_limit):
        page = doc[page_index]
        images = page.get_images(full=True)
        if len(images) > MAX_IMAGE_CANDIDATES_PER_PAGE:
            skipped += len(images)
            continue
        for image_index, img in enumerate(images):
            xref = img[0]
            try:
                image_info = doc.extract_image(xref)
            except Exception:
                skipped += 1
                continue
            image_bytes = image_info.get("image", b"") or b""
            width = int(image_info.get("width") or 0)
            height = int(image_info.get("height") or 0)
            if width < MIN_IMAGE_EDGE_PX or height < MIN_IMAGE_EDGE_PX or len(image_bytes) < MIN_IMAGE_BYTES:
                skipped += 1
                continue
            caption = find_caption(text_lines_by_page.get(page_index + 1, []), image_index)
            text = caption or f"[Image page {page_index + 1} #{image_index + 1}]"
            blocks.append(
                Block(
                    type="image",
                    text=text,
                    metadata={
                        "page": page_index + 1,
                        "image_index": image_index,
                        "width": width,
                        "height": height,
                        "asset_size_bytes": len(image_bytes),
                        "asset_sha256": hashlib.sha256(image_bytes).hexdigest(),
                    },
                )
            )
    if skipped:
        meta["skipped_image_candidates"] = skipped


def ocr_pdf_pages(doc: Any, page_limit: int) -> list[Block]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:
        return []
    out: list[Block] = []
    for page_index in range(page_limit):
        page = doc[page_index]
        pix = page.get_pixmap(matrix=getattr(__import__("fitz"), "Matrix")(2, 2), alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = normalize_space(pytesseract.image_to_string(img, lang="chi_sim+eng"))
        if text:
            out.append(Block(type="paragraph", text=text, metadata={"page": page_index + 1, "ocr": True}))
    return out


def parse_docx(data: bytes) -> dict[str, Any]:
    try:
        import docx  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is required for DOCX parsing") from exc

    document = docx.Document(io.BytesIO(data))
    blocks: list[Block] = []
    heading_stack: list[str] = []
    for idx, paragraph in enumerate(document.paragraphs):
        text = normalize_space(paragraph.text)
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        level = docx_heading_level(style_name)
        if level:
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(text)
            blocks.append(
                Block(
                    type="heading",
                    text=text,
                    heading_path=list(heading_stack),
                    heading_inferred=False,
                    heading_confidence=1.0,
                    metadata={"paragraph_index": idx, "style": style_name},
                )
            )
        elif looks_like_formula(text):
            blocks.append(Block(type="formula", text=text, heading_path=list(heading_stack), metadata={"paragraph_index": idx}))
        else:
            blocks.append(Block(type="paragraph", text=text, heading_path=list(heading_stack), metadata={"paragraph_index": idx}))

    for table_index, table in enumerate(document.tables):
        rows = [[normalize_space(cell.text) for cell in row.cells] for row in table.rows]
        md = table_to_markdown(rows)
        if md:
            blocks.append(Block(type="table", text=md, metadata={"table_index": table_index}))

    return {
        "meta": {"parser": "python-docx", "paragraph_count": len(document.paragraphs), "table_count": len(document.tables)},
        "blocks": [b.as_json() for b in blocks],
    }


def parse_image(data: bytes, mime: str) -> dict[str, Any]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Pillow and pytesseract are required for image parsing") from exc

    image = Image.open(io.BytesIO(data))
    width, height = image.size
    text = normalize_space(pytesseract.image_to_string(image, lang="chi_sim+eng"))
    blocks: list[Block] = []
    metadata = {
        "width": width,
        "height": height,
        "asset_size_bytes": len(data),
        "asset_sha256": hashlib.sha256(data).hexdigest(),
        "ocr": True,
    }
    if width >= MIN_IMAGE_EDGE_PX and height >= MIN_IMAGE_EDGE_PX and len(data) >= MIN_IMAGE_BYTES:
        blocks.append(Block(type="image", text=text or "[Image OCR produced no text]", metadata=metadata))
    elif text:
        blocks.append(Block(type="paragraph", text=text, metadata={**metadata, "image_filtered": True}))
    return {"meta": {"parser": "pillow+tesseract", "image_mime": mime}, "blocks": [b.as_json() for b in blocks]}


def finalize(payload: dict[str, Any]) -> None:
    blocks = payload.get("blocks") or []
    counts = Counter(str(block.get("type", "paragraph")) for block in blocks)
    inferred = sum(1 for block in blocks if block.get("type") == "heading" and block.get("heading_inferred"))
    headings = sum(1 for block in blocks if block.get("type") == "heading")
    meta = payload.setdefault("meta", {})
    meta["block_type_counts"] = dict(counts)
    meta["heading_inferred_ratio"] = (inferred / headings) if headings else 0.0


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def weighted_font_size(spans: list[dict[str, Any]]) -> float:
    total = sum(max(1, len(s.get("text", ""))) for s in spans)
    if total <= 0:
        return 0.0
    return sum(float(s.get("size") or 0) * max(1, len(s.get("text", ""))) for s in spans) / total


def dominant_font(spans: list[dict[str, Any]]) -> str:
    counts: Counter[str] = Counter()
    for span in spans:
        font = str(span.get("font") or "")
        counts[font] += max(1, len(span.get("text", "")))
    return counts.most_common(1)[0][0] if counts else ""


def is_heading_line(line: dict[str, Any], body_size: float) -> bool:
    text = str(line["text"])
    size = float(line["font_size"])
    font = str(line.get("font") or "").lower()
    short_enough = len(text) <= 80
    numbered = bool(re.match(r"^(\d+(\.\d+)*|第[一二三四五六七八九十百]+[章节])", text))
    return short_enough and (size >= body_size + 1.2 or "bold" in font or numbered)


def heading_confidence(line: dict[str, Any], body_size: float) -> float:
    size = float(line["font_size"])
    delta = max(0.0, size - body_size)
    return min(0.95, 0.55 + delta / 10.0)


def looks_like_formula(text: str) -> bool:
    if len(text) > 160:
        return False
    formula_chars = sum(1 for ch in text if ch in "=+-*/^√∑∫≤≥≈△Δλμ")
    return formula_chars >= 2 and bool(re.search(r"[A-Za-z0-9\u0394\u03bb\u03bc]", text))


def table_to_markdown(table: list[list[Any]]) -> str:
    rows = [[normalize_space("" if cell is None else str(cell)) for cell in row] for row in table if row]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = padded[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def find_caption(lines: list[str], image_index: int) -> str:
    caption_re = re.compile(r"^(图|表|Figure|Fig\.)\s*[\d一二三四五六七八九十\-.]*", re.IGNORECASE)
    matches = [line for line in lines if caption_re.search(line)]
    if not matches:
        return ""
    return matches[min(image_index, len(matches) - 1)]


def median(values: list[float]) -> float:
    if not values:
        return 0.0
    sorted_values = sorted(values)
    mid = len(sorted_values) // 2
    if len(sorted_values) % 2:
        return sorted_values[mid]
    return (sorted_values[mid - 1] + sorted_values[mid]) / 2.0


def docx_heading_level(style_name: str) -> int:
    match = re.search(r"heading\s*(\d+)", style_name, re.IGNORECASE)
    if match:
        return max(1, min(6, int(match.group(1))))
    match = re.search(r"标题\s*(\d+)", style_name)
    if match:
        return max(1, min(6, int(match.group(1))))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
